"""OEM Datasheet Rebrander routes"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from typing import Optional
from datetime import datetime, timezone
from uuid import uuid4
import logging
import os
import io
import re
import json
import zipfile
import tempfile

from middleware.auth import get_current_user, is_admin
from db.connection import get_db

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/oem", tags=["oem"])

UPLOAD_DIR = os.environ.get("UPLOAD_DIR", "/app/uploads")
BRAND_LOGOS_DIR = os.path.join(UPLOAD_DIR, "brand_logos")
os.makedirs(BRAND_LOGOS_DIR, exist_ok=True)


def _resolve_logo_path(filename: str, brand: dict) -> str | None:
    """Return a usable file path for a brand logo.

    1. Try the normal filesystem path.
    2. If missing (e.g. after redeploy), restore the file from the base64
       stored in MongoDB and return the restored path.
    3. Return None if the logo cannot be found or restored.
    """
    import base64 as _b64
    lpath = os.path.join(BRAND_LOGOS_DIR, filename)
    if os.path.exists(lpath):
        return lpath
    # Try to restore from MongoDB base64
    logo_data_map = brand.get("logoDataMap") or {}
    b64 = logo_data_map.get(filename)
    if not b64:
        return None
    try:
        os.makedirs(BRAND_LOGOS_DIR, exist_ok=True)
        with open(lpath, "wb") as f:
            f.write(_b64.b64decode(b64))
        return lpath
    except Exception:
        return None

# Colors considered neutral — never replaced with brand color
NEUTRAL_COLORS = {
    'FFFFFF', 'FEFEFE', 'FDFDFD', 'F9F9F9', 'F8F8F8',
    '000000', '000001', '010101', '1A1A1A', '0D0D0D',
    'F2F2F2', 'F0F0F0', 'EBEBEB', 'E6E6E6', 'E5E5E5', 'DEDEDE',
    'D9D9D9', 'CCCCCC', 'C0C0C0', 'BFBFBF', 'B2B2B2', 'A6A6A6',
    '808080', '737373', '666666', '595959', '4D4D4D', '404040',
    '333333', '262626', '1F1F1F',
    'AUTO',
}


# ==================== HELPERS ====================

def extract_text_from_docx(content: bytes) -> str:
    """Extract all text from a Word document"""
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF using pdfplumber or PyPDF2 as fallback"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"pdfplumber failed: {e}")
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception as e2:
            logger.error(f"PyPDF2 also failed: {e2}")
            return ""


def extract_images_from_pdf(pdf_content: bytes) -> list:
    import fitz
    doc = fitz.open(stream=pdf_content, filetype="pdf")
    images = []
    for page_num, page in enumerate(doc):
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            images.append({
                "page": page_num,
                "index": img_index,
                "ext": base_image["ext"],
                "data": base_image["image"]
            })
    return images


def convert_pdf_to_docx(pdf_content: bytes) -> bytes:
    """
    Convert PDF → DOCX while preserving layout (tables, columns, fonts).
    Uses pdf2docx which reconstructs the document structure.
    """
    try:
        from pdf2docx import Converter
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="pdf2docx library not installed. Cannot preserve PDF layout."
        )

    pdf_tmp = None
    docx_tmp = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(pdf_content)
            pdf_tmp = f.name

        docx_tmp = pdf_tmp.replace(".pdf", ".docx")
        cv = Converter(pdf_tmp)
        cv.convert(docx_tmp, start=0, end=None)
        cv.close()

        with open(docx_tmp, "rb") as f:
            return f.read()
    finally:
        if pdf_tmp and os.path.exists(pdf_tmp):
            os.unlink(pdf_tmp)
        if docx_tmp and os.path.exists(docx_tmp):
            os.unlink(docx_tmp)


def replace_text_in_docx(content: bytes, replacements: dict) -> bytes:
    """
    Replace supplier text with brand text.
    Operates run-by-run to preserve character formatting.
    """
    from docx import Document

    if not replacements:
        return content

    doc = Document(io.BytesIO(content))

    def apply(text: str) -> str:
        for old, new in replacements.items():
            if old and old.strip() and new is not None:
                text = text.replace(old, new)
        return text

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                if run.text:
                    run.text = apply(run.text)
            # Fix cross-run splits
            full = "".join(r.text for r in para.runs)
            if apply(full) != full:
                for run in para.runs:
                    run.text = apply(run.text)

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)
    for section in doc.sections:
        process_paragraphs(section.header.paragraphs)
        process_paragraphs(section.footer.paragraphs)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def replace_images_in_docx(content: bytes, brand_logo_paths: list) -> bytes:
    """
    Replace existing inline images in a docx with brand logos (ZIP-level).
    Keeps original image positions and sizes intact.
    """
    if not brand_logo_paths:
        return content

    with zipfile.ZipFile(io.BytesIO(content), "r") as zin:
        zip_data = {name: zin.read(name) for name in zin.namelist()}

    media_files = sorted(
        [n for n in zip_data if n.startswith("word/media/") and not n.endswith("/")]
    )

    if not media_files:
        return insert_logos_at_top(content, brand_logo_paths)

    for i, media_key in enumerate(media_files):
        logo_path = brand_logo_paths[i % len(brand_logo_paths)]
        if not os.path.exists(logo_path):
            continue
        with open(logo_path, "rb") as f:
            logo_bytes = f.read()
        media_ext = os.path.splitext(media_key)[1].lower()
        logo_ext = os.path.splitext(logo_path)[1].lower()
        if logo_ext == media_ext:
            zip_data[media_key] = logo_bytes
        else:
            try:
                from PIL import Image
                img = Image.open(io.BytesIO(logo_bytes))
                buf = io.BytesIO()
                if media_ext in (".jpg", ".jpeg"):
                    img = img.convert("RGB")
                    img.save(buf, "JPEG")
                elif media_ext == ".gif":
                    img.save(buf, "GIF")
                else:
                    img.save(buf, "PNG")
                zip_data[media_key] = buf.getvalue()
            except Exception as e:
                logger.warning(f"Logo format conversion failed: {e}")
                zip_data[media_key] = logo_bytes

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in zip_data.items():
            zout.writestr(name, data)
    return out.getvalue()


def replace_colors_in_docx(content: bytes, primary_color: str, secondary_color: str = None) -> bytes:
    """
    Replace non-neutral colors in DOCX XML with brand colors (ZIP-level).
    Targets:
      - w:val="RRGGBB"     → text colors
      - w:fill="RRGGBB"    → paragraph / cell shading
      - a:srgbClr val=...  → DrawingML shape / table-header colors
      - a:lastClr val=...  → DrawingML theme gradient colors
    """
    primary = primary_color.lstrip('#').upper() if primary_color else ''
    secondary = secondary_color.lstrip('#').upper() if secondary_color else primary

    if not primary:
        return content

    with zipfile.ZipFile(io.BytesIO(content), "r") as zin:
        zip_data = {name: zin.read(name) for name in zin.namelist()}

    xml_files = [
        n for n in zip_data
        if n.startswith("word/") and n.endswith(".xml")
    ]

    for xml_file in xml_files:
        try:
            text = zip_data[xml_file].decode("utf-8")
        except Exception:
            continue

        original = text

        def sub_w_val(m):
            c = m.group(1).upper()
            return m.group(0) if c in NEUTRAL_COLORS else f'w:val="{primary}"'

        def sub_w_fill(m):
            c = m.group(1).upper()
            return m.group(0) if c in NEUTRAL_COLORS else f'w:fill="{secondary}"'

        def sub_srgb(m):
            c = m.group(1).upper()
            return m.group(0) if c in NEUTRAL_COLORS else f'val="{primary}"'

        text = re.sub(r'w:val="([0-9A-Fa-f]{6})"', sub_w_val, text)
        text = re.sub(r'w:fill="([0-9A-Fa-f]{6})"', sub_w_fill, text)
        text = re.sub(r'(?<=<a:srgbClr )val="([0-9A-Fa-f]{6})"', sub_srgb, text)
        text = re.sub(r'(?<=<a:lastClr )val="([0-9A-Fa-f]{6})"', sub_srgb, text)

        if text != original:
            zip_data[xml_file] = text.encode("utf-8")

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in zip_data.items():
            zout.writestr(name, data)
    return out.getvalue()


def insert_logos_at_top(content: bytes, logo_paths: list) -> bytes:
    """Fallback: insert brand logos at the top when document has no existing images."""
    from docx import Document
    from docx.shared import Inches

    doc = Document(io.BytesIO(content))
    for logo_path in reversed(logo_paths):
        if not os.path.exists(logo_path):
            continue
        try:
            first = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
            run = first.insert_paragraph_before().add_run()
            run.add_picture(logo_path, width=Inches(1.5))
        except Exception as e:
            logger.warning(f"Could not insert logo {logo_path}: {e}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def identify_supplier_info(text: str, openai_api_key: str) -> dict:
    """Ask GPT-4o-mini to identify supplier brand info in the document text"""
    from openai import AsyncOpenAI
    client = AsyncOpenAI(api_key=openai_api_key)

    prompt = f"""Analyze this product datasheet text and extract the supplier/manufacturer brand information.

Return ONLY a JSON object with these fields (use null if not found):
{{
  "company_name": "...",
  "address": "...",
  "phone": "...",
  "email": "...",
  "website": "...",
  "logo_alt_text": "...",
  "other_brand_mentions": ["...", "..."]
}}

Datasheet text (first 3000 chars):
{text[:3000]}"""

    response = await client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"},
        temperature=0
    )
    try:
        return json.loads(response.choices[0].message.content)
    except Exception:
        return {}


async def rebuild_docx_from_pdf(
    pdf_content: bytes,
    brand: dict,
    supplier_info: dict,
    openai_api_key: str,
    extracted_images: list = None,
) -> bytes:
    """
    AI-powered PDF → clean DOCX rebuild:
    1. Extract text from PDF
    2. GPT parses it into structured sections + tables (JSON)
    3. python-docx builds a clean DOCX with brand colors, logo in header,
       brand contact in footer. Images are omitted (user adds manually).
    """
    from openai import AsyncOpenAI
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    text = extract_text_from_pdf(pdf_content)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from PDF")

    # Build supplier → brand replacement map for post-processing
    replacements = {}
    if supplier_info.get("company_name") and brand.get("name"):
        replacements[supplier_info["company_name"]] = brand["name"]
    if supplier_info.get("address") and brand.get("address"):
        replacements[supplier_info["address"]] = brand["address"]
    if supplier_info.get("phone") and brand.get("phone"):
        replacements[supplier_info["phone"]] = brand["phone"]
    if supplier_info.get("email") and brand.get("email"):
        replacements[supplier_info["email"]] = brand["email"]
    if supplier_info.get("website") and brand.get("website"):
        replacements[supplier_info["website"]] = brand["website"]
    for mention in (supplier_info.get("other_brand_mentions") or []):
        if mention and mention.strip():
            replacements[mention] = brand["name"]

    def apply_replacements(s: str) -> str:
        if not s:
            return s
        for old, new in replacements.items():
            if old:
                s = s.replace(old, new)
        return s

    # GPT: parse document structure
    client = AsyncOpenAI(api_key=openai_api_key)
    prompt = f"""Parse this product datasheet text into a structured JSON.

Return ONLY valid JSON with exactly this structure:
{{
  "title": "product model/name (e.g. GYXTS)",
  "sections": [
    {{
      "heading": "section title including number (e.g. '1. Cable cross-section')",
      "text": "plain text description, or null",
      "table": {{
        "headers": ["Column1", "Column2", ...],
        "rows": [["val1", "val2", ...], ...]
      }}
    }}
  ]
}}

Rules:
- Include ALL sections and ALL table data with exact values
- Skip any images or diagrams (they will be added manually)
- Do NOT include supplier company name, logo, address, phone, email or website — those will be replaced
- Preserve technical specifications exactly (numbers, units, symbols)
- If a section has both text and a table, include both fields

Datasheet text:
{text[:5000]}"""

    try:
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0
        )
        structure = json.loads(response.choices[0].message.content)
    except Exception as e:
        logger.error(f"GPT structure parsing failed: {e}")
        structure = {"title": text[:80].split("\n")[0], "sections": []}

    # Brand colors
    primary_hex = (brand.get("primaryColor") or "#1E3A8A").lstrip("#")
    subtitle_hex = (brand.get("subtitleColor") or brand.get("primaryColor") or "#1E3A8A").lstrip("#")

    def hex_to_rgb(h):
        h = (h or "1E3A8A").lstrip("#")
        try:
            return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
        except Exception:
            return (30, 58, 138)

    primary_rgb = hex_to_rgb(primary_hex)
    subtitle_rgb = hex_to_rgb(subtitle_hex)

    # px → units conversion (96dpi screen)
    PX_TO_IN   = 1 / 96
    PX_TO_PT   = 72 / 96
    PX_TO_TWIP = 1440 / 96   # 1 inch = 1440 twips

    header_height_px  = int(brand.get("headerHeightPx")  or 60)
    header_padding_px = int(brand.get("headerPaddingPx") or 8)
    footer_height_px  = int(brand.get("footerHeightPx")  or 36)
    footer_padding_px = int(brand.get("footerPaddingPx") or 6)
    logo_size_px      = int(brand.get("logoSizePx")      or max(10, header_height_px - 2 * header_padding_px))

    copyright_text = brand.get("copyrightText") or f"All rights reserved © {datetime.now(timezone.utc).year}"

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
        tcPr.append(shd)

    def _set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
        tcPr.append(shd)

    def _set_table_no_borders(tbl_elem):
        """Remove all visible borders from a table."""
        tblPr = tbl_elem.tblPr if tbl_elem.tblPr is not None else OxmlElement("w:tblPr")
        tblBdr = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
            tblBdr.append(b)
        tblPr.append(tblBdr)

    def _set_row_exact_height(row, height_px):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trH = OxmlElement("w:trHeight")
        trH.set(qn("w:val"), str(int(height_px * PX_TO_TWIP)))
        trH.set(qn("w:hRule"), "exact")
        trPr.append(trH)

    def _set_cell_padding(cell, px):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        pt_val = str(int(px * PX_TO_TWIP))
        for side in ("top", "left", "bottom", "right"):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"), pt_val)
            m.set(qn("w:type"), "dxa")
            tcMar.append(m)
        tcPr.append(tcMar)

    def _cell_no_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBdr = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tcBdr.append(b)
        tcPr.append(tcBdr)

    # Build DOCX
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Inches(header_height_px * PX_TO_IN + 0.3)
    section.bottom_margin = Inches(footer_height_px * PX_TO_IN + 0.2)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.header_distance = Inches(0.1)
    section.footer_distance = Inches(0.1)

    PAGE_MARGIN_TWIPS    = int(1.0 * 1440)   # matches section.left_margin = Inches(1.0)
    FOOTER_EMAIL_LEFT_PX = 10

    def _apply_full_width_band(para, hex_color, height_px, padding_px, content_left_px):
        """Full-width colored band (edge-to-edge).

        Technique: negative w:ind left/right extends the paragraph shading into
        page margins.  w:firstLine positions the content (logo/text) at the
        desired offset from the left MARGIN.

        Band = before(padding) + exact-line(inner) + after(padding)
        Content starts at: content_left_px from the left margin.
        """
        inner_px = max(4, height_px - 2 * padding_px)
        pPr = para._p.get_or_add_pPr()

        # Full-width background fill
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
        pPr.append(shd)

        # Vertical: padding top+bottom, exact inner height
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"),   str(int(padding_px * PX_TO_TWIP)))
        spacing.set(qn("w:after"),    str(int(padding_px * PX_TO_TWIP)))
        spacing.set(qn("w:line"),     str(int(inner_px   * PX_TO_TWIP)))
        spacing.set(qn("w:lineRule"), "exact")
        pPr.append(spacing)

        # Horizontal: extend band to paper edges; firstLine puts content at desired position
        content_twips = int(content_left_px * PX_TO_TWIP)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"),      str(-PAGE_MARGIN_TWIPS))
        ind.set(qn("w:right"),     str(-PAGE_MARGIN_TWIPS))
        # firstLine is relative to w:left → content_left_px from LEFT MARGIN
        ind.set(qn("w:firstLine"), str(PAGE_MARGIN_TWIPS + content_twips))
        pPr.append(ind)

    # ── Header: full-width colored band, logo left ─────────────────────
    header = section.header
    header.is_linked_to_previous = False

    h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    h_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _apply_full_width_band(h_para, primary_hex, header_height_px,
                           padding_px=header_padding_px,
                           content_left_px=header_padding_px)

    logo_added = False
    for logo_fn in (brand.get("approvedLogos") or []):
        lpath = _resolve_logo_path(logo_fn, brand)
        if lpath:
            try:
                inner_px  = max(4, header_height_px - 2 * header_padding_px)
                capped_px = min(logo_size_px, inner_px)
                h_para.add_run().add_picture(lpath, height=Inches(max(0.1, capped_px * PX_TO_IN)))
                logo_added = True
                break
            except Exception:
                pass
    if not logo_added:
        r = h_para.add_run(brand.get("name", ""))
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    header_image_fn = brand.get("headerImage")
    if header_image_fn:
        hi_path = _resolve_logo_path(header_image_fn, brand)
        if hi_path:
            try:
                inner_px = max(4, header_height_px - 2 * header_padding_px)
                h_para.add_run().add_picture(hi_path, height=Inches(max(0.1, inner_px * PX_TO_IN)))
            except Exception as e:
                logger.warning(f"Could not insert header image: {e}")

    # ── Footer: full-width colored band, email left | copyright center ─
    footer = section.footer
    footer.is_linked_to_previous = False

    f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    _apply_full_width_band(f_para, primary_hex, footer_height_px,
                           padding_px=footer_padding_px,
                           content_left_px=FOOTER_EMAIL_LEFT_PX)

    # Center tab stop — measured from left margin (not paper edge)
    pPr = f_para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    ct = OxmlElement("w:tab")
    ct.set(qn("w:val"), "center")
    ct.set(qn("w:pos"), "4680")
    tabs.append(ct)
    pPr.append(tabs)

    r_email = f_para.add_run(brand.get("email") or brand.get("name") or "")
    r_email.font.size = Pt(8)
    r_email.font.color.rgb = RGBColor(255, 255, 255)

    f_para.add_run("\t")

    r_copy = f_para.add_run(copyright_text)
    r_copy.font.size = Pt(8)
    r_copy.font.color.rgb = RGBColor(255, 255, 255)

    # Document title
    title_text = apply_replacements(structure.get("title", ""))
    if title_text:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(title_text)
        tr.bold = True
        tr.font.size = Pt(16)
        tr.font.color.rgb = RGBColor(*primary_rgb)

    doc.add_paragraph()

    # Sections
    for sec_index, sec in enumerate(structure.get("sections", [])):
        heading = apply_replacements(sec.get("heading") or "")
        text_body = apply_replacements(sec.get("text") or "")
        table_data = sec.get("table")

        if heading:
            hp = doc.add_paragraph()
            hr = hp.add_run(heading)
            hr.bold = True
            hr.font.size = Pt(11)
            hr.font.color.rgb = RGBColor(*subtitle_rgb)

        if text_body:
            bp = doc.add_paragraph(text_body)
            for run in bp.runs:
                run.font.size = Pt(10)

        if table_data:
            headers = table_data.get("headers") or []
            rows = table_data.get("rows") or []
            col_count = max(
                len(headers),
                max((len(r) for r in rows), default=0)
            )
            if col_count > 0 and (headers or rows):
                tbl = doc.add_table(rows=len(rows) + (1 if headers else 0), cols=col_count)
                tbl.style = "Table Grid"

                row_offset = 0
                if headers:
                    hdr_row = tbl.rows[0].cells
                    for ci, h in enumerate(headers[:col_count]):
                        cell = hdr_row[ci]
                        set_cell_bg(cell, primary_hex)
                        r = cell.paragraphs[0].add_run(apply_replacements(str(h)))
                        r.bold = True
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(255, 255, 255)
                    row_offset = 1

                for ri, row in enumerate(rows):
                    cells = tbl.rows[ri + row_offset].cells
                    for ci, val in enumerate(row[:col_count]):
                        r = cells[ci].paragraphs[0].add_run(apply_replacements(str(val)))
                        r.font.size = Pt(9)

        doc.add_paragraph()

        # Insert extracted PDF image for this section (skip index 0 — likely supplier logo)
        img_insert_index = sec_index + 1
        if extracted_images and img_insert_index < len(extracted_images):
            try:
                img_buf = io.BytesIO(extracted_images[img_insert_index]["data"])
                doc.add_picture(img_buf, width=Inches(4.0))
                doc.add_paragraph()
            except Exception as e:
                logger.warning(f"Could not insert PDF image {img_insert_index}: {e}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ==================== LOGO SERVE ENDPOINT ====================

@router.get("/logo/{filename}")
async def serve_logo(filename: str):
    """Serve brand logo files"""
    from fastapi.responses import FileResponse
    if ".." in filename or "/" in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    db = get_db()
    # Find the brand that owns this logo
    brand = await db.oem_brands.find_one({"approvedLogos": filename}, {"_id": 0})
    logo_path = _resolve_logo_path(filename, brand or {})
    if not logo_path:
        raise HTTPException(status_code=404, detail="Logo not found")
    return FileResponse(logo_path)


# ==================== BRAND ENDPOINTS ====================

@router.get("/brands")
async def list_brands(current_user: dict = Depends(get_current_user)):
    db = get_db()
    brands = await db.oem_brands.find({}, {"_id": 0}).to_list(100)
    return brands


@router.post("/brands")
async def create_brand(
    name: str = Form(...),
    address: str = Form(""),
    phone: str = Form(""),
    email: str = Form(""),
    website: str = Form(""),
    warrantyText: str = Form(""),
    primaryColor: str = Form(""),
    subtitleColor: str = Form(""),
    headerHeightPx: int = Form(60),
    headerPaddingPx: int = Form(8),
    logoSizePx: int = Form(44),
    footerHeightPx: int = Form(36),
    footerPaddingPx: int = Form(6),
    copyrightText: str = Form(""),
    current_user: dict = Depends(get_current_user)
):
    if not is_admin(current_user["email"]):
        raise HTTPException(status_code=403, detail="Admin only")
    db = get_db()
    brand = {
        "id": str(uuid4()),
        "name": name,
        "address": address,
        "phone": phone,
        "email": email,
        "website": website,
        "warrantyText": warrantyText,
        "primaryColor": primaryColor,
        "subtitleColor": subtitleColor,
        "headerHeightPx": headerHeightPx,
        "headerPaddingPx": headerPaddingPx,
        "logoSizePx": logoSizePx,
        "footerHeightPx": footerHeightPx,
        "footerPaddingPx": footerPaddingPx,
        "copyrightText": copyrightText,
        "approvedLogos": [],
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    await db.oem_brands.insert_one({**brand, "_id": brand["id"]})
    return brand


@router.put("/brands/{brand_id}")
async def update_brand(
    brand_id: str,
    name: str = Form(...),
    address: str = Form(""),
    phone: str = Form(""),
    email: str = Form(""),
    website: str = Form(""),
    warrantyText: str = Form(""),
    primaryColor: str = Form(""),
    subtitleColor: str = Form(""),
    headerHeightPx: int = Form(60),
    headerPaddingPx: int = Form(8),
    logoSizePx: int = Form(44),
    footerHeightPx: int = Form(36),
    footerPaddingPx: int = Form(6),
    copyrightText: str = Form(""),
    current_user: dict = Depends(get_current_user)
):
    if not is_admin(current_user["email"]):
        raise HTTPException(status_code=403, detail="Admin only")
    db = get_db()
    update = {
        "name": name,
        "address": address,
        "phone": phone,
        "email": email,
        "website": website,
        "warrantyText": warrantyText,
        "primaryColor": primaryColor,
        "subtitleColor": subtitleColor,
        "headerHeightPx": headerHeightPx,
        "headerPaddingPx": headerPaddingPx,
        "logoSizePx": logoSizePx,
        "footerHeightPx": footerHeightPx,
        "footerPaddingPx": footerPaddingPx,
        "copyrightText": copyrightText,
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    result = await db.oem_brands.update_one({"id": brand_id}, {"$set": update})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Brand not found")
    brand = await db.oem_brands.find_one({"id": brand_id}, {"_id": 0})
    return brand


@router.delete("/brands/{brand_id}")
async def delete_brand(brand_id: str, current_user: dict = Depends(get_current_user)):
    if not is_admin(current_user["email"]):
        raise HTTPException(status_code=403, detail="Admin only")
    db = get_db()
    result = await db.oem_brands.delete_one({"id": brand_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Brand not found")
    return {"success": True}


@router.post("/brands/{brand_id}/logo")
async def upload_brand_logo(
    brand_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    if not is_admin(current_user["email"]):
        raise HTTPException(status_code=403, detail="Admin only")
    db = get_db()
    brand = await db.oem_brands.find_one({"id": brand_id})
    if not brand:
        raise HTTPException(status_code=404, detail="Brand not found")

    allowed = {".png", ".jpg", ".jpeg", ".svg", ".webp"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed:
        raise HTTPException(status_code=400, detail="Only PNG, JPG, SVG, WEBP allowed")

    filename = f"{brand_id}_{str(uuid4())[:8]}{ext}"
    logo_path = os.path.join(BRAND_LOGOS_DIR, filename)
    content = await file.read()
    with open(logo_path, "wb") as f:
        f.write(content)

    # Store base64 in MongoDB so logo survives server redeployments
    import base64 as _b64
    logo_b64 = _b64.b64encode(content).decode("utf-8")
    logo_data_map = brand.get("logoDataMap") or {}
    logo_data_map[filename] = logo_b64

    approved_logos = brand.get("approvedLogos", [])
    approved_logos.append(filename)
    await db.oem_brands.update_one(
        {"id": brand_id},
        {"$set": {
            "approvedLogos": approved_logos,
            "logoDataMap": logo_data_map,
            "updatedAt": datetime.now(timezone.utc).isoformat()
        }}
    )
    return {"filename": filename, "approvedLogos": approved_logos}


@router.delete("/brands/{brand_id}/logo/{filename}")
async def delete_brand_logo(
    brand_id: str,
    filename: str,
    current_user: dict = Depends(get_current_user)
):
    if not is_admin(current_user["email"]):
        raise HTTPException(status_code=403, detail="Admin only")
    db = get_db()
    brand = await db.oem_brands.find_one({"id": brand_id})
    if not brand:
        raise HTTPException(status_code=404, detail="Brand not found")

    logo_path = os.path.join(BRAND_LOGOS_DIR, filename)
    if os.path.exists(logo_path):
        os.remove(logo_path)

    approved_logos = [l for l in (brand.get("approvedLogos") or []) if l != filename]
    logo_data_map = brand.get("logoDataMap") or {}
    logo_data_map.pop(filename, None)
    await db.oem_brands.update_one(
        {"id": brand_id},
        {"$set": {"approvedLogos": approved_logos, "logoDataMap": logo_data_map}}
    )
    return {"approvedLogos": approved_logos}


@router.post("/brands/{brand_id}/header-image")
async def upload_brand_header_image(
    brand_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    if not is_admin(current_user["email"]):
        raise HTTPException(status_code=403, detail="Admin only")
    db = get_db()
    brand = await db.oem_brands.find_one({"id": brand_id})
    if not brand:
        raise HTTPException(status_code=404, detail="Brand not found")
    allowed = {".png", ".jpg", ".jpeg", ".svg", ".webp"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed:
        raise HTTPException(status_code=400, detail="Only PNG, JPG, SVG, WEBP allowed")
    filename = f"header_{brand_id}_{str(uuid4())[:8]}{ext}"
    path = os.path.join(BRAND_LOGOS_DIR, filename)
    content = await file.read()
    with open(path, "wb") as f:
        f.write(content)
    import base64 as _b64
    logo_data_map = brand.get("logoDataMap") or {}
    logo_data_map[filename] = _b64.b64encode(content).decode("utf-8")
    await db.oem_brands.update_one(
        {"id": brand_id},
        {"$set": {
            "headerImage": filename,
            "logoDataMap": logo_data_map,
            "updatedAt": datetime.now(timezone.utc).isoformat()
        }}
    )
    return {"filename": filename}


# ==================== DATASHEET PROCESSING ====================

@router.post("/process")
async def process_datasheet(
    file: UploadFile = File(...),
    brand_id: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    """
    Upload supplier datasheet → AI identifies supplier info → rebrand → return OEM DOCX.

    PDF  → AI parses structure → clean DOCX rebuild (brand colors, logo, tables)
    DOCX → in-place text/image/color replacement (layout fully preserved)
    """
    db = get_db()

    brand = await db.oem_brands.find_one({"id": brand_id}, {"_id": 0})
    if not brand:
        raise HTTPException(status_code=404, detail="Brand not found")

    raw_content = await file.read()
    orig_filename = file.filename
    filename_lower = orig_filename.lower()

    if not (filename_lower.endswith(".pdf") or filename_lower.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only .docx and .pdf files are supported")

    is_pdf = filename_lower.endswith(".pdf")

    openai_key = os.environ.get("OPENAI_API_KEY", "")

    # --- Identify supplier info (used for both PDF and DOCX paths) ---
    supplier_info = {}
    if openai_key:
        try:
            raw_text = extract_text_from_pdf(raw_content) if is_pdf else extract_text_from_docx(raw_content)
            supplier_info = await identify_supplier_info(raw_text, openai_key)
            logger.info(f"Supplier info identified: {supplier_info}")
        except Exception as e:
            logger.error(f"Supplier identification failed: {e}")

    safe_brand = re.sub(r"[^a-zA-Z0-9_-]", "_", brand["name"])
    base_name = os.path.splitext(orig_filename)[0]
    extracted_images = []

    if is_pdf:
        # ── PDF: AI rebuild → clean DOCX with brand colors + logo ──────
        if not openai_key:
            raise HTTPException(status_code=500, detail="OPENAI_API_KEY not configured")
        try:
            extracted_images = extract_images_from_pdf(raw_content)
            logger.info(f"Extracted {len(extracted_images)} images from PDF")
        except Exception as e:
            logger.warning(f"Image extraction failed: {e}")
        output_bytes = await rebuild_docx_from_pdf(raw_content, brand, supplier_info, openai_key, extracted_images)
        output_filename = f"OEM_{safe_brand}_{base_name}.docx"
    else:
        # ── DOCX: in-place text + image + color replacement ─────────────
        replacements = {}
        if supplier_info.get("company_name") and brand.get("name"):
            replacements[supplier_info["company_name"]] = brand["name"]
        if supplier_info.get("address") and brand.get("address"):
            replacements[supplier_info["address"]] = brand["address"]
        if supplier_info.get("phone") and brand.get("phone"):
            replacements[supplier_info["phone"]] = brand["phone"]
        if supplier_info.get("email") and brand.get("email"):
            replacements[supplier_info["email"]] = brand["email"]
        if supplier_info.get("website") and brand.get("website"):
            replacements[supplier_info["website"]] = brand["website"]
        for mention in (supplier_info.get("other_brand_mentions") or []):
            if mention and mention.strip():
                replacements[mention] = brand["name"]

        output_bytes = replace_text_in_docx(raw_content, replacements)

        logo_paths = [
            p for p in (
                _resolve_logo_path(fn, brand)
                for fn in (brand.get("approvedLogos") or [])
            ) if p
        ]
        if logo_paths:
            output_bytes = replace_images_in_docx(output_bytes, logo_paths)

        primary_color = brand.get("primaryColor", "")
        if primary_color:
            output_bytes = replace_colors_in_docx(output_bytes, primary_color)

        output_filename = f"OEM_{safe_brand}_{base_name}.docx"

    # --- Log job ---
    await db.oem_jobs.insert_one({
        "_id": str(uuid4()),
        "userId": current_user["id"],
        "brandId": brand_id,
        "brandName": brand["name"],
        "originalFilename": orig_filename,
        "supplierInfo": supplier_info,
        "outputFormat": "pdf→docx" if is_pdf else "docx",
        "extractedImageCount": len(extracted_images) if is_pdf else 0,
        "createdAt": datetime.now(timezone.utc).isoformat(),
    })

    return StreamingResponse(
        io.BytesIO(output_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{output_filename}"'}
    )
