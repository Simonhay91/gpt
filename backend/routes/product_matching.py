"""Product Matching — upload customer file, AI-match against catalog, preview + Excel."""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Query
from bson import ObjectId
from services.planet_api import get_catalog as _planet_get_catalog, get_product_datasheet_url as _get_datasheet_url
from fastapi.responses import FileResponse
from starlette.background import BackgroundTask
from typing import List, Optional
import io
import os
import re
import json
import uuid
import logging
from datetime import datetime, timezone
from pydantic import BaseModel

import numpy as np
import anthropic
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

from db.connection import get_db
from middleware.auth import get_current_user

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/product-matching", tags=["product-matching"])

CLAUDE_API_KEY = os.environ.get("CLAUDE_API_KEY", "")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
MAX_FILE_SIZE = 20 * 1024 * 1024   # 20 MB
MAX_CUSTOMER_ITEMS = 200
MAX_CATALOG_PRODUCTS = 5000        # load more — Voyage pre-filters, Claude never sees all
VOYAGE_TOP_K = 10                  # candidates per item from Voyage similarity
PHASE3_TOP_K = 12                  # candidates after web-research enrichment
WEB_SEARCH_ENABLED = True          # Phase 3: web research fallback for unmatched items
CLAUDE_BATCH_SIZE = 30             # customer items per Claude call (Phase 2)


# ==================== TEMPLATE BUILDER ====================

# Keywords that identify the "product name" column in customer Excel files
_PRODUCT_COL_KEYWORDS = (
    "product", "name", "item", "description", "наименование",
    "товар", "продукт", "անվանում", "ապրանք", "model", "модель",
)

_SKIP_VALUES = {"nan", "none", "", "qty", "quantity", "price", "notes", "note",
                "колич", "цена", "примечание", "remarks", "კომენტარი"}


def _build_template_excel() -> bytes:
    """Generate the customer template Excel file and return as bytes."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Products"

    header_fill = PatternFill("solid", fgColor="FFF2CC")   # soft yellow
    header_font = Font(bold=True, size=12, color="1F4E79")
    hint_font   = Font(italic=True, size=10, color="808080")
    border_side = openpyxl.styles.Side(style="thin", color="BFBFBF")
    border      = openpyxl.styles.Border(
        left=border_side, right=border_side,
        top=border_side, bottom=border_side,
    )

    headers = ["Product Name", "Qty", "Notes"]
    col_widths = [45, 10, 30]

    for col_idx, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.fill   = header_fill
        cell.font   = header_font
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = w

    ws.row_dimensions[1].height = 22

    # Hint row
    hint_cell = ws.cell(row=2, column=1, value="← Paste your product names here (required)")
    hint_cell.font = hint_font
    hint_cell.alignment = Alignment(vertical="center")

    ws.cell(row=2, column=2, value="").font = hint_font
    ws.cell(row=2, column=3, value="").font = hint_font

    # Freeze header row
    ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ==================== FILE PARSERS ====================

def _parse_excel(content: bytes) -> List[str]:
    import pandas as pd
    buf = io.BytesIO(content)

    # Try reading with header row first
    try:
        df_with_header = pd.read_excel(buf, header=0)
    except Exception:
        buf.seek(0)
        try:
            df_with_header = pd.read_csv(buf, header=0, on_bad_lines="skip")
        except Exception:
            df_with_header = None

    # Detect product name column by header keyword
    target_col = None
    if df_with_header is not None:
        for col in df_with_header.columns:
            col_lower = str(col).lower().strip()
            if any(kw in col_lower for kw in _PRODUCT_COL_KEYWORDS):
                target_col = col
                break

    if target_col is not None:
        # Use the detected named column (skip header — already excluded by header=0)
        series = df_with_header[target_col].dropna()
    else:
        # Fallback: use first column, read without header to include all rows
        buf.seek(0)
        try:
            df_no_header = pd.read_excel(buf, header=None)
        except Exception:
            buf.seek(0)
            try:
                df_no_header = pd.read_csv(buf, header=None, on_bad_lines="skip")
            except Exception:
                buf.seek(0)
                df_no_header = pd.read_csv(buf, on_bad_lines="skip")
        series = df_no_header.iloc[:, 0].dropna()

    seen: set = set()
    items: List[str] = []
    for val in series:
        s = str(val).strip()
        if s and s.lower() not in _SKIP_VALUES and s not in seen:
            seen.add(s)
            items.append(s)
    return items


def _parse_docx(content: bytes) -> List[str]:
    from docx import Document
    doc = Document(io.BytesIO(content))
    seen: set = set()
    items: List[str] = []

    def _add(text: str):
        s = text.strip()
        if s and s not in seen:
            seen.add(s)
            items.append(s)

    for para in doc.paragraphs:
        _add(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _add(cell.text)
    return items


def _parse_pdf(content: bytes) -> List[str]:
    import pdfplumber
    seen: set = set()
    items: List[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for line in text.splitlines():
                s = line.strip()
                if s and s not in seen:
                    seen.add(s)
                    items.append(s)
    return items


# ==================== CLAUDE MATCHING ====================

def _cosine_similarity(a: list, b: list) -> float:
    va, vb = np.array(a, dtype=np.float32), np.array(b, dtype=np.float32)
    denom = np.linalg.norm(va) * np.linalg.norm(vb)
    return float(np.dot(va, vb) / denom) if denom > 0 else 0.0


def _voyage_embed_batch(texts: List[str]) -> List[Optional[List[float]]]:
    """Embed a list of texts via OpenAI text-embedding-3-small."""
    from openai import OpenAI
    client = OpenAI(api_key=OPENAI_API_KEY)
    response = client.embeddings.create(
        input=[t[:8000] for t in texts],
        model="text-embedding-3-small",
    )
    return [item.embedding for item in response.data]


def _text_fallback_top_k(item: str, catalog: List[dict], k: int) -> List[dict]:
    """Simple case-insensitive text match fallback when embeddings are unavailable."""
    item_lower = item.lower()
    hits = []
    for p in catalog:
        haystack = " ".join(filter(None, [
            p.get("title_en", ""),
            p.get("article_number", ""),
            " ".join(p.get("aliases") or []),
        ])).lower()
        if item_lower in haystack or any(w in haystack for w in item_lower.split() if len(w) > 3):
            hits.append(p)
        if len(hits) >= k:
            break
    return hits[:k]


def _web_research_item(item: str, client: anthropic.Anthropic) -> tuple:
    """
    Use Claude with the built-in web_search tool to research an unknown product.
    Returns (enriched_description: str, source_urls: List[str]).
    """
    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1024,
            tools=[{"type": "web_search_20250305", "name": "web_search"}],
            messages=[{
                "role": "user",
                "content": (
                    f'Research this product: "{item}". '
                    f"Find its full technical name, manufacturer, product category, "
                    f"key specifications (fiber count, construction type, armor, sheath, standard). "
                    f"Return a concise technical product description in English (2-4 sentences)."
                ),
            }],
        )

        enriched = ""
        source_urls: List[str] = []

        for block in response.content:
            if block.type == "text":
                enriched = block.text.strip()
            elif block.type == "tool_result":
                # web_search returns list of results with url field
                for entry in (block.content or []):
                    url = (entry.get("url") or "") if isinstance(entry, dict) else ""
                    if url and url not in source_urls:
                        source_urls.append(url)

        # Also check tool_use blocks for search result metadata
        for block in response.content:
            if hasattr(block, "type") and block.type == "tool_use":
                pass  # search query visible here but URLs come from tool_result

        if not enriched:
            enriched = item  # fallback to original if Claude returned nothing

        return enriched, source_urls[:5]

    except Exception as exc:
        logger.warning(f"Web research failed for '{item}': {exc}")
        return item, []


def _phase3_web_rematch(
    unmatched_items: List[tuple],   # [(orig_idx, customer_item), ...]
    catalog: List[dict],
    catalog_embeddings: List[Optional[List[float]]],
) -> dict:
    """
    Phase 3: for each unmatched item, run web research then re-match against catalog.
    Returns dict: orig_idx → result dict (with web_sources field).
    """
    if not unmatched_items:
        return {}

    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
    results: dict = {}

    # Research each item individually (web search is sequential by nature)
    enriched_data: List[tuple] = []   # (orig_idx, item, enriched_text, source_urls)
    for orig_idx, item in unmatched_items:
        logger.info(f"Phase 3: web research for '{item}'")
        enriched_text, source_urls = _web_research_item(item, client)
        enriched_data.append((orig_idx, item, enriched_text, source_urls))

    # Re-embed enriched descriptions and find new candidates
    enriched_texts = [enriched for _, _, enriched, _ in enriched_data]
    try:
        enriched_candidates = _voyage_top_k(
            enriched_texts, catalog, catalog_embeddings, PHASE3_TOP_K
        )
    except Exception as exc:
        logger.warning(f"Phase 3 Voyage re-embed failed: {exc}")
        enriched_candidates = [
            _text_fallback_top_k(enriched, catalog, PHASE3_TOP_K)
            for enriched in enriched_texts
        ]

    # Claude re-match with enriched context
    items_with_candidates = [
        (item, candidates)
        for (_, item, _, _), candidates in zip(enriched_data, enriched_candidates)
    ]

    try:
        claude_results = _claude_match_with_candidates(items_with_candidates)
    except Exception as exc:
        logger.error(f"Phase 3 Claude re-match failed: {exc}")
        claude_results = _empty_batch([item for item, _ in items_with_candidates], "Web research failed")

    for i, (orig_idx, item, enriched_text, source_urls) in enumerate(enriched_data):
        r = claude_results[i] if i < len(claude_results) else {
            "customer_item": item,
            "matched_title": "", "article_number": "",
            "crm_code": "", "vendor": "", "datasheet_url": "",
            "confidence": "none", "comment": None,
        }
        r["web_sources"] = source_urls
        results[orig_idx] = r

    return results


def _voyage_top_k(
    items: List[str],
    catalog: List[dict],
    catalog_embeddings: List[Optional[List[float]]],
    k: int,
) -> List[List[dict]]:
    """
    For each item return TOP-k catalog products by cosine similarity.
    Falls back to text search for items/products without embeddings.
    """
    try:
        item_embeddings = _voyage_embed_batch(items)
    except Exception as exc:
        logger.warning(f"Voyage embed error for items: {exc}")
        return [_text_fallback_top_k(item, catalog, k) for item in items]

    results = []
    for item, item_emb in zip(items, item_embeddings):
        if not item_emb:
            results.append(_text_fallback_top_k(item, catalog, k))
            continue

        scores = []
        for i, (p, p_emb) in enumerate(zip(catalog, catalog_embeddings)):
            if p_emb:
                score = _cosine_similarity(item_emb, p_emb)
            else:
                # no embedding for this product → text fallback score
                haystack = " ".join(filter(None, [
                    p.get("title_en", ""), p.get("article_number", ""),
                ])).lower()
                score = 0.3 if item.lower() in haystack else 0.0
            scores.append((score, i))

        top_indices = [i for _, i in sorted(scores, reverse=True)[:k]]
        results.append([catalog[i] for i in top_indices])

    return results


def _format_candidates(candidates: List[dict]) -> str:
    """Format a small list of catalog candidates for the Claude prompt."""
    lines = []
    for i, p in enumerate(candidates):
        aliases = ", ".join((p.get("aliases") or [])[:3])
        parts = [
            f"  [{i + 1}] title={p.get('title_en', '')}",
            f"article={p.get('article_number', '')}",
            f"crm={p.get('crm_code', '')}",
            f"vendor={p.get('vendor', '')}",
        ]
        if p.get("product_model"):
            parts.append(f"model={p['product_model']}")
        if aliases:
            parts.append(f"aliases={aliases}")
        lines.append(" | ".join(parts))
    return "\n".join(lines) if lines else "  (no candidates found)"


def _claude_match_with_candidates(
    items_with_candidates: List[tuple],  # [(customer_item, [candidate_dicts]), ...]
    domain: str = "",
) -> List[dict]:
    """
    Send ALL items + their TOP-k candidates to Claude in ONE request.
    Returns list of match dicts in same order.
    """
    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

    items_text_parts = []
    for i, (item, candidates) in enumerate(items_with_candidates):
        cand_text = _format_candidates(candidates)
        items_text_parts.append(
            f"[{i + 1}] Customer item: \"{item}\"\n"
            f"     Candidates:\n{cand_text}"
        )
    items_text = "\n\n".join(items_text_parts)

    effective_domain = domain if domain else OPTICAL_CABLE_DOMAIN
    prompt = f"""You are a product matching assistant for a fiber optics and network equipment catalog.

{effective_domain}

For each customer item below I have pre-selected the TOP candidates from the catalog using semantic search.
Choose the best matching candidate, or return no match if none fits.

{items_text}

Return a JSON array with exactly {len(items_with_candidates)} objects in the same order:
[
  {{
    "customer_item": "<original customer item text>",
    "matched_title": "<chosen candidate title_en, empty string if no match>",
    "article_number": "<chosen candidate article_number, empty string if no match>",
    "crm_code": "<chosen candidate crm_code, empty string if no match>",
    "vendor": "<chosen candidate vendor, empty string if no match>",
    "datasheet_url": "",
    "confidence": "<high|medium|low|none>",
    "comment": "<always required: one short English sentence explaining the match or why no match was found>"
  }}
]

Rules:
- confidence="high": very strong match, customer item clearly maps to the candidate
- confidence="medium": reasonable match but some ambiguity (different spec, slightly different model)
- confidence="low": weak match, best available but uncertain
- confidence="none": no suitable match among candidates
- If confidence is "none", set matched_title/article_number/crm_code to empty strings
- comment is ALWAYS required (never null or empty):
  • high   → e.g. "Exact match: same model code and fiber count"
  • medium → e.g. "Approximate: customer requested 24FO, matched 48FO variant"
  • low    → e.g. "Weak match: closest available but different construction type"
  • none   → e.g. "No match: GYTA 6-module/8-core variant not found in catalog"
- Return ONLY the JSON array, no extra text."""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = response.content[0].text.strip()
    m = re.search(r"\[.*\]", raw, re.DOTALL)
    if not m:
        logger.error(f"Claude (candidates) returned no JSON array: {raw[:300]}")
        return _empty_batch([item for item, _ in items_with_candidates], "Matching error")
    try:
        return json.loads(m.group(0))
    except json.JSONDecodeError as exc:
        logger.error(f"Claude (candidates) JSON decode error: {exc} | raw[:300]={raw[:300]}")
        return _empty_batch([item for item, _ in items_with_candidates], "Matching error")


def _format_catalog(catalog: List[dict]) -> str:
    lines = []
    for i, p in enumerate(catalog):
        aliases = ", ".join((p.get("aliases") or [])[:5])  # max 5 aliases to save tokens
        parts = [
            f"[{i + 1}]",
            f"title={p.get('title_en', '')}",
            f"article={p.get('article_number', '')}",
            f"crm={p.get('crm_code', '')}",
            f"vendor={p.get('vendor', '')}",
        ]
        if p.get('product_model'):
            parts.append(f"model={p['product_model']}")
        if aliases:
            parts.append(f"aliases={aliases}")
        lines.append(" | ".join(parts))
    return "\n".join(lines)


OPTICAL_CABLE_DOMAIN = """
OPTICAL FIBER CABLE NAMING CONVENTION (use this to interpret customer requests):

MODEL FORMAT: [PREFIX]-[TUBE]([MODIFIER])-([FIBER_RANGE])FO([SPEC])-[STRENGTH]KN

--- PREFIX (installation type) ---
AS       = Aerial ADSS self-supporting
AM       = Aerial Messenger / Drop outdoor (Fig-8 or flat)
U-TIC    = Underground Steel Tape armored
U-DIC    = Underground Steel Wire armored
U-TBC    = Underground Air Blow Cable
OM3      = Multimode OM3
OM4      = Multimode OM4
IR       = Indoor Riser
ID       = Indoor Drop

--- TUBE CONSTRUCTION ---
L   = Central Loose Tube
M   = Loose Tube + Central Strength member (mix)
S   = Multi Tube / Central Strength member
ST  = Single Tube (indoor)
MT  = Multi Tube (indoor)
FL  = Flat Drop cable
OR  = O-ring / Fig-8 shape

--- MODIFIERS (in parentheses after tube type) ---
(A)    = Aramid yarns reinforcement
(PA)   = Partial Aramid yarns
(D)    = FRP messenger
(W)    = Steel wire strength member
(L)    = Double sheath (e.g. AS(L)-S = double-sheath ADSS)
(KFRP) = Kevlar FRP reinforcement

--- FIBER COUNT ---
Written as (min-max)FO, e.g.:
  (1-24)FO  = supports 1 to 24 fibers
  (4-72)FO  = supports 4 to 72 fibers
  (96-144)FO = supports 96 to 144 fibers
  (2-480)FO = supports 2 to 480 fibers
Match customer's fiber count to the range that contains it.

--- FIBER SPEC SUFFIX ---
(none)   = Standard SMF ITU-T G.652D
LSZH     = Low Smoke Zero Halogen
G657A2   = Bend-insensitive ITU-T G.657A2 (indoor/drop)
Aqua     = OM3 indoor aqua color
300m/1km = Drum length

--- TENSILE STRENGTH (trailing KN) ---
Common values: 0.5, 0.6, 0.8, 1, 1.5, 2, 2.7, 2.8, 3, 3.5, 4, 4.5, 5.5, 7, 7.5, 8, 9, 9.5, 20, 25 KN

--- GYTA / GYTS / GYFTY CABLES (ITU-T standard naming) ---
G  = Gel-filled loose tubes
Y  = Polyethylene outer sheath
T  = Steel tape armor
A  = Aluminum moisture barrier
S  = Steel wire strength member (GYTS = steel wire instead of tape)

GYTA structure: central strength member + loose tubes (each tube holds 6–12 fibers)
  "X modules / Y core" = X tubes × Y fibers per tube → total = X×Y fibers
  e.g. "6 modules/8 core" = 6 tubes × 8 fibers = 48 fibers total → match to 48-fiber GYTA
       "4 modules/6 core" = 4 tubes × 6 fibers = 24 fibers total → match to 24-fiber GYTA
       "2 modules/8 core" = 2 tubes × 8 fibers = 16 fibers total → match to 16-fiber GYTA
Outer diameter clue: 10.5mm ≈ 48FO, 9.5mm ≈ 16–24FO (use as tie-breaker only)
When searching catalog for GYTA: use total fiber count as primary key.

--- MATCHING EXAMPLES ---
Customer says → Best model match:
"ADSS 12 fiber 2KN"                    → AS-M-(4-12)FO-2KN      (12 falls in 4-12 range)
"ADSS 48fo 3kn strength member"        → AS-S-(4-72)FO-3KN       (48 in 4-72, strength member)
"aerial loose tube 24 fiber 1.5kn"     → AS-L-(1-24)FO-1.5KN
"ADSS double sheath 96fo 20kn aramid"  → AS(L)-S-(8-96)FO-20KN
"fig-8 loose tube 12fo 5.5kn"         → AM-L-(1-24)FO-5.5KN
"drop flat FRP messenger 2fo 0.6kn"   → AM-L-FL(D)-(1-4)FO-0.6KN
"drop o-ring aramid 4fo 1kn"          → AM-L-OR(A)-(1-4)FO-1KN
"underground steel tape 24fo 2.7kn"   → U-TIC-L-(2-24)FO-2.7KN
"underground steel wire multi 48fo"   → U-DIC-S-(8-96)FO-7KN
"air blow cable 96 fiber"             → U-TBC-S-(2-480)FO
"indoor riser 48fo single tube"       → IR-ST-(8-24)FO-G657A2
"indoor riser multi tube 72fo"        → IR-MT-(14-96)FO-G657A2
"indoor drop flat FRP 2fo"            → ID-FL(D)-(1-4)FO-G657A2
"indoor drop flat steel wire 4fo"     → ID-FL(W)-(1-4)FO-G657A2
"indoor drop o-ring 300m drum"        → ID-OR3MM-(1-4)FO-G657A2(300m)
"om3 multimode 12fo 1kn lszh"         → OM3-L-(6-12)FO-1kN-LSZH
"om4 multimode cable"                 → OM4-L-(6-12)FO-1kN-LSZH
"drop adss aramid 12fo 1kn"           → AS-L(A)-(4-24)FO-1KN
"""


def _claude_match_batch(
    batch: List[str],
    catalog_text: str,
    catalog_len: int,
    domain: str = "",
) -> List[dict]:
    """
    Send one batch of customer items + full catalog to Claude.
    Returns list of match dicts in the same order as batch.
    """
    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
    items_text = "\n".join(f"[{i + 1}] {name}" for i, name in enumerate(batch))

    effective_domain = domain if domain else OPTICAL_CABLE_DOMAIN
    prompt = f"""You are a product matching assistant for a fiber optics and network equipment catalog.

{effective_domain}

CATALOG ({catalog_len} products):
{catalog_text}

CUSTOMER ITEMS ({len(batch)} items to match):
{items_text}

For each customer item find the best matching catalog product.
Use the optical cable naming convention above to interpret both customer requests and catalog model codes.
Return a JSON array with exactly {len(batch)} objects in the same order:
[
  {{
    "customer_item": "<original customer item text>",
    "matched_title": "<catalog title_en, empty string if no match>",
    "article_number": "<article_number from catalog, empty string if no match>",
    "crm_code": "<crm_code from catalog, empty string if no match>",
    "vendor": "<vendor from catalog, empty string if no match>",
    "datasheet_url": "<datasheet_url from catalog, empty string if no match>",
    "comment": "<always required: one short English sentence explaining the match or why no match was found>"
  }}
]

Rules:
- For optical cables: decode the customer's description using the naming convention, then find the catalog model whose fiber range contains the requested count and whose strength is closest.
- comment is ALWAYS required (never null or empty):
  • exact match   → e.g. "Exact match: same model and fiber count"
  • approximate   → e.g. "Approximate: customer 1.5KN, matched 2KN variant"
  • no match      → e.g. "No match: requested fiber count not available in catalog"
- Return ONLY the JSON array, no extra text."""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = response.content[0].text.strip()
    m = re.search(r"\[.*\]", raw, re.DOTALL)
    if not m:
        logger.error(f"Claude returned no JSON array: {raw[:300]}")
        return _empty_batch(batch, "Matching error")
    try:
        return json.loads(m.group(0))
    except json.JSONDecodeError as exc:
        logger.error(f"Claude JSON decode error: {exc} | raw[:300]={raw[:300]}")
        return _empty_batch(batch, "Matching error")


def _empty_batch(items: List[str], comment: str) -> List[dict]:
    return [
        {
            "customer_item": item,
            "matched_title": "",
            "article_number": "",
            "crm_code": "",
            "vendor": "",
            "datasheet_url": "",
            "comment": comment,
        }
        for item in items
    ]


# ==================== EXCEL BUILDER ====================

def _build_excel(rows: List[dict], mode: str) -> str:
    """Build result Excel, save to /tmp/, return path."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Product Matching"

    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    alt_fill = PatternFill("solid", fgColor="EBF3FB")
    unmatch_fill = PatternFill("solid", fgColor="FFF2CC")

    code_label = "CRM Code" if mode == "global" else "Article Number"
    headers = ["Customer Item", code_label, "Our Product Title", "Datasheet URL", "Comment"]

    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 22

    col_widths = [35, 20, 40, 40, 45]
    for col_idx, w in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = w

    for row_idx, r in enumerate(rows, 2):
        is_matched = bool(r.get("matched_title", "").strip())
        is_alt_row = (row_idx % 2 == 0)

        if not is_matched:
            fill = unmatch_fill
        elif is_alt_row:
            fill = alt_fill
        else:
            fill = None

        values = [
            r.get("customer_item", ""),
            r.get("code", ""),
            r.get("matched_title", ""),
            r.get("datasheet_url", ""),
            r.get("comment") or "",
        ]

        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            if fill:
                cell.fill = fill
            cell.alignment = Alignment(
                vertical="center",
                wrap_text=(col_idx in (1, 5)),
            )

    path = f"/tmp/product_matching_{uuid.uuid4()}.xlsx"
    wb.save(path)
    return path


def _remove_file(path: str):
    try:
        os.unlink(path)
    except Exception:
        pass


# ==================== PYDANTIC MODELS ====================

class GenerateRequest(BaseModel):
    mode: str = "global"
    results: List[dict]


class ResearchItemRequest(BaseModel):
    item: str
    mode: str = "global"


class PlanetSearchRequest(BaseModel):
    query: str


class DomainRuleCreate(BaseModel):
    title: str
    content: str
    category: str = "general"   # "vendor_naming" | "cable_type" | "general"
    is_active: bool = True


class DomainRuleUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    category: Optional[str] = None
    is_active: Optional[bool] = None


# ==================== ENDPOINTS ====================

PLANET_SEARCH_TOP_K = 5   # max CRM codes returned to PlanetWorkspace


# ── PlanetWorkspace fallback search ──────────────────────────────────────────

@router.post("/planet-search")
async def planet_search(data: PlanetSearchRequest):
    """
    Called by PlanetWorkspace when a user search returns no results.
    Runs alias lookup → Voyage AI pre-filter → Claude match against catalog.
    Returns an ordered array of CRM codes (best match first, max 5).
    """

    query = (data.query or "").strip()
    if not query:
        raise HTTPException(status_code=400, detail="query is required")

    db = get_db()

    catalog = await _planet_get_catalog(db)

    if not catalog:
        raise HTTPException(status_code=404, detail="Product catalog is empty")

    # ── Step 1: Alias lookup (instant, no AI) ────────────────────────────────
    alias_lookup: dict = {}
    for p in catalog:
        for alias in (p.get("aliases") or []):
            alias_lookup[alias.lower().strip()] = p

    crm_to_product = {p.get("crm_code", ""): p for p in catalog if p.get("crm_code")}
    article_to_product = {p.get("article_number", ""): p for p in catalog if p.get("article_number")}
    saved_aliases = await db.product_aliases.find({}, {"_id": 0}).to_list(10000)
    for sa in saved_aliases:
        key = (sa.get("alias") or "").lower().strip()
        if key and key not in alias_lookup:
            prod = (
                crm_to_product.get(sa.get("crm_code") or "")
                or article_to_product.get(sa.get("article_number") or "")
            )
            if prod:
                alias_lookup[key] = prod

    query_lower = query.lower().strip()
    if query_lower in alias_lookup:
        p = alias_lookup[query_lower]
        crm = p.get("crm_code")
        logger.info(f"planet-search alias hit: '{query}' → {crm}")
        return [crm] if crm else []

    # ── Step 2: Voyage AI pre-filter → Claude match ───────────────────────────
    catalog_embeddings = [p.get("embedding") for p in catalog]

    try:
        candidates_list = _voyage_top_k([query], catalog, catalog_embeddings, VOYAGE_TOP_K)
        candidates = candidates_list[0] if candidates_list else []
    except Exception as exc:
        logger.warning(f"planet-search Voyage failed, using text fallback: {exc}")
        candidates = _text_fallback_top_k(query, catalog, VOYAGE_TOP_K)

    if not candidates:
        logger.info(f"planet-search: no candidates found for '{query}'")
        return []

    try:
        claude_results = _claude_match_with_candidates([(query, candidates)])
        r = claude_results[0] if claude_results else {}
    except Exception as exc:
        logger.error(f"planet-search Claude failed: {exc}")
        return []

    confidence = (r.get("confidence") or "").lower()
    crm = r.get("crm_code") or ""

    if confidence == "none" or not crm:
        logger.info(f"planet-search: no match for '{query}' (confidence={confidence})")
        return []

    logger.info(f"planet-search: '{query}' → {crm} (confidence={confidence})")
    return [crm]


# ── Domain Rules CRUD ────────────────────────────────────────────────────────

@router.get("/domain-rules")
async def list_domain_rules(current_user: dict = Depends(get_current_user)):
    """Return all matching domain rules."""
    db = get_db()
    rules = await db.matching_domain_rules.find(
        {}, {"embedding": 0}
    ).sort("updated_at", -1).to_list(200)
    for r in rules:
        r["_id"] = str(r["_id"])
    return rules


@router.post("/domain-rules", status_code=201)
async def create_domain_rule(
    body: DomainRuleCreate,
    current_user: dict = Depends(get_current_user),
):
    """Create a new matching domain rule."""
    db = get_db()
    doc = {
        **body.model_dump(),
        "created_by": str(current_user.get("_id", "")),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    result = await db.matching_domain_rules.insert_one(doc)
    doc["_id"] = str(result.inserted_id)
    return doc


@router.put("/domain-rules/{rule_id}")
async def update_domain_rule(
    rule_id: str,
    body: DomainRuleUpdate,
    current_user: dict = Depends(get_current_user),
):
    """Update an existing matching domain rule."""
    from bson import ObjectId
    db = get_db()
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="No fields to update")
    updates["updated_at"] = datetime.now(timezone.utc).isoformat()
    result = await db.matching_domain_rules.update_one(
        {"_id": ObjectId(rule_id)}, {"$set": updates}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Rule not found")
    rule = await db.matching_domain_rules.find_one({"_id": ObjectId(rule_id)})
    rule["_id"] = str(rule["_id"])
    return rule


@router.delete("/domain-rules/{rule_id}", status_code=204)
async def delete_domain_rule(
    rule_id: str,
    current_user: dict = Depends(get_current_user),
):
    """Delete a matching domain rule."""
    from bson import ObjectId
    db = get_db()
    result = await db.matching_domain_rules.delete_one({"_id": ObjectId(rule_id)})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Rule not found")

@router.get("/template")
async def download_template(current_user: dict = Depends(get_current_user)):
    """Return a pre-formatted customer Excel template for product matching."""
    template_bytes = _build_template_excel()
    path = f"/tmp/product_matching_template_{uuid.uuid4()}.xlsx"
    with open(path, "wb") as f:
        f.write(template_bytes)
    return FileResponse(
        path=path,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename="product_matching_template.xlsx",
        background=BackgroundTask(_remove_file, path),
    )


@router.post("/match")
async def match_products(
    file: UploadFile = File(...),
    mode: str = Form("global"),   # "global" | "oem"
    current_user: dict = Depends(get_current_user),
):
    """
    Upload a customer file (xlsx/xls/csv/docx/pdf) with product names.
    Returns JSON preview results for user review before Excel generation.

    3-phase pipeline:
      Step 0  — alias lookup (instant, no AI)
      Phase 1 — Voyage AI embedding pre-filter → TOP-5 candidates per item
      Phase 2 — Claude picks best candidate from the small set
    Fallback: if Voyage AI unavailable → full catalog sent to Claude (original logic).
    """
    db = get_db()

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large (max 20 MB)")

    filename = (file.filename or "").lower()

    # 1. Parse customer items from file
    try:
        if filename.endswith(".pdf"):
            customer_items = _parse_pdf(content)
        elif filename.endswith(".docx"):
            customer_items = _parse_docx(content)
        else:
            customer_items = _parse_excel(content)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {exc}")

    if not customer_items:
        raise HTTPException(status_code=400, detail="No product names found in file")

    customer_items = customer_items[:MAX_CUSTOMER_ITEMS]

    # 2. Load catalog from PlanetWorkspace API (normalized + embeddings)
    catalog = await _planet_get_catalog(db)

    if not catalog:
        raise HTTPException(status_code=404, detail="Product catalog is empty")

    # 3. Load OEM vendor names (for OEM mode code selection)
    oem_vendor_names: set = set()
    if mode == "oem":
        oem_brands = await db.oem_brands.find({}, {"_id": 0, "name": 1}).to_list(200)
        oem_vendor_names = {b["name"].lower() for b in oem_brands if b.get("name")}

    # 4. Load active custom domain rules and build full domain context
    custom_rules = await db.matching_domain_rules.find(
        {"is_active": True}, {"_id": 0, "content": 1}
    ).to_list(100)
    if custom_rules:
        custom_domain = "\n\n".join(r["content"] for r in custom_rules if r.get("content"))
        full_domain = OPTICAL_CABLE_DOMAIN + "\n\n--- CUSTOM VENDOR RULES ---\n" + custom_domain
    else:
        full_domain = OPTICAL_CABLE_DOMAIN

    # ── STEP 0: Alias lookup ─────────────────────────────────────────────────
    # Build lookup from catalog inline aliases
    alias_lookup: dict = {}
    for p in catalog:
        for alias in (p.get("aliases") or []):
            alias_lookup[alias.lower().strip()] = p

    # Also load saved product_aliases collection
    crm_to_product = {p.get("crm_code", ""): p for p in catalog if p.get("crm_code")}
    article_to_product = {p.get("article_number", ""): p for p in catalog if p.get("article_number")}
    saved_aliases = await db.product_aliases.find({}, {"_id": 0}).to_list(10000)
    for sa in saved_aliases:
        key = (sa.get("alias") or "").lower().strip()
        if key and key not in alias_lookup:
            prod = (
                crm_to_product.get(sa.get("crm_code") or "")
                or article_to_product.get(sa.get("article_number") or "")
            )
            if prod:
                alias_lookup[key] = prod

    alias_matched: dict = {}   # orig_idx → result dict
    needs_ai: List[tuple] = [] # (orig_idx, item)

    for idx, item in enumerate(customer_items):
        item_lower = item.lower().strip()
        if item_lower in alias_lookup:
            p = alias_lookup[item_lower]
            alias_matched[idx] = {
                "customer_item": item,
                "matched_title": p.get("title_en", ""),
                "article_number": p.get("article_number", ""),
                "crm_code": p.get("crm_code", ""),
                "vendor": p.get("vendor", ""),
                "datasheet_url": p.get("datasheet_url", ""),
                "comment": None,
                "match_type": "auto",
                "confidence": "high",
            }
        else:
            needs_ai.append((idx, item))

    logger.info(
        f"Product matching: total={len(customer_items)}, "
        f"alias_hit={len(alias_matched)}, needs_ai={len(needs_ai)}"
    )

    # ── PHASE 1 + 2: Voyage pre-filter → Claude ──────────────────────────────
    ai_map: dict = {}  # orig_idx → raw result dict

    if needs_ai:
        embed_ok = bool(OPENAI_API_KEY)
        ai_items = [item for _, item in needs_ai]
        ai_indices = [idx for idx, _ in needs_ai]

        if embed_ok:
            # ── Phase 1: OpenAI embedding → TOP-K candidates ──────────────────
            try:
                catalog_embeddings = [p.get("embedding") for p in catalog]
                items_candidates = _voyage_top_k(
                    ai_items, catalog, catalog_embeddings, VOYAGE_TOP_K
                )
                logger.info(
                    f"Phase 1 (OpenAI embed) complete: {len(ai_items)} items → "
                    f"{VOYAGE_TOP_K} candidates each"
                )
                use_candidates = True
            except Exception as exc:
                logger.warning(f"Embedding phase failed, falling back to full-catalog: {exc}")
                use_candidates = False
        else:
            logger.warning("OPENAI_API_KEY not set — falling back to full-catalog Claude")
            use_candidates = False

        if use_candidates:
            # ── Phase 2: Claude with small candidate lists ────────────────────
            items_with_candidates = list(zip(ai_items, items_candidates))
            raw_claude: List[dict] = []
            for i in range(0, len(items_with_candidates), CLAUDE_BATCH_SIZE):
                batch = items_with_candidates[i: i + CLAUDE_BATCH_SIZE]
                try:
                    batch_results = _claude_match_with_candidates(batch, domain=full_domain)
                except Exception as exc:
                    logger.error(f"Claude (candidates) batch error: {exc}")
                    batch_results = _empty_batch([item for item, _ in batch], "Matching error")
                raw_claude.extend(batch_results)
            logger.info(f"Phase 2 (Claude with candidates) complete: {len(raw_claude)} results")
        else:
            # ── Fallback: full catalog → Claude (original logic) ──────────────
            catalog_text = _format_catalog(catalog)
            raw_claude = []
            for i in range(0, len(ai_items), CLAUDE_BATCH_SIZE):
                batch = ai_items[i: i + CLAUDE_BATCH_SIZE]
                try:
                    batch_results = _claude_match_batch(batch, catalog_text, len(catalog), domain=full_domain)
                except Exception as exc:
                    logger.error(f"Claude (full-catalog) batch error: {exc}")
                    batch_results = _empty_batch(batch, "Matching error")
                raw_claude.extend(batch_results)
            logger.info(f"Fallback (full-catalog Claude) complete: {len(raw_claude)} results")

        for claude_idx, r in enumerate(raw_claude):
            orig_idx = ai_indices[claude_idx]
            ai_map[orig_idx] = r

    # ── Enrich datasheet_url from catalog (Claude always returns it empty) ────
    catalog_by_crm = {p["crm_code"]: p for p in catalog if p.get("crm_code")}
    catalog_by_article = {p["article_number"]: p for p in catalog if p.get("article_number")}
    for r in ai_map.values():
        if not r.get("datasheet_url"):
            matched_p = catalog_by_crm.get(r.get("crm_code") or "") or \
                        catalog_by_article.get(r.get("article_number") or "")
            if matched_p and matched_p.get("datasheet_url"):
                r["datasheet_url"] = matched_p["datasheet_url"]

    # ── Merge & build final response ─────────────────────────────────────────
    results: List[dict] = []
    for idx, item in enumerate(customer_items):
        if idx in alias_matched:
            r = alias_matched[idx]
        else:
            r = ai_map.get(idx, {
                "customer_item": item,
                "matched_title": "", "article_number": "",
                "crm_code": "", "vendor": "", "datasheet_url": "",
                "comment": "Matching error",
            })

            # Normalise confidence from Claude response
            raw_conf = (r.get("confidence") or "").lower()
            if raw_conf in ("high", "medium", "low"):
                confidence = raw_conf
            else:
                has_match = bool(r.get("matched_title", "").strip())
                confidence = ("medium" if r.get("comment") else "high") if has_match else None
            r["match_type"] = "auto"
            r["confidence"] = confidence

        # Apply mode logic to pick code
        vendor_lower = (r.get("vendor") or "").lower()
        article = r.get("article_number") or ""
        crm = r.get("crm_code") or ""

        if mode == "global":
            code = crm or article
        else:
            if vendor_lower and vendor_lower in oem_vendor_names:
                code = article or crm
            else:
                code = crm or article

        results.append({
            "customer_item": r.get("customer_item", ""),
            "crm_code": crm or None,
            "article_number": article or None,
            "matched_title": r.get("matched_title", ""),
            "code": code or None,
            "datasheet_url": r.get("datasheet_url", "") or None,
            "comment": r.get("comment") or None,
            "web_sources": [],
            "match_type": r.get("match_type", "auto"),
            "confidence": r.get("confidence"),
        })

    # ── Enrich datasheet_url with PDF links from product detail ─────────────
    # catalog_by_crm/article already built above; catalog products have 'slug'
    needs_pdf: List[tuple] = []  # (result_idx, slug)
    for i, r in enumerate(results):
        if r.get("matched_title") and "/api/public/" not in (r.get("datasheet_url") or ""):
            crm = r.get("crm_code") or ""
            article = r.get("article_number") or ""
            matched_p = catalog_by_crm.get(crm) or catalog_by_article.get(article)
            if matched_p and matched_p.get("slug"):
                needs_pdf.append((i, matched_p["slug"]))

    if needs_pdf:
        import asyncio as _asyncio
        pdf_urls = await _asyncio.gather(
            *[_get_datasheet_url(slug) for _, slug in needs_pdf],
            return_exceptions=True,
        )
        for (idx, _), url in zip(needs_pdf, pdf_urls):
            if isinstance(url, str) and url:
                results[idx]["datasheet_url"] = url
        logger.info(f"Datasheet enrichment: fetched {len(needs_pdf)} URLs")

    logger.info(f"Product matching done: {len(results)} results returned")
    return {"results": results}


@router.post("/research-item")
async def research_item(
    data: ResearchItemRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Web-research a single unmatched customer item, then re-match against catalog.
    Called manually from the preview UI when the user clicks 'Research' on a row.
    Returns a single match result dict (same shape as /match results[i]).
    """
    db = get_db()
    item = (data.item or "").strip()
    mode = data.mode

    if not item:
        raise HTTPException(status_code=400, detail="item is required")

    # Load OEM vendor names if needed
    oem_vendor_names: set = set()
    if mode == "oem":
        oem_brands = await db.oem_brands.find({}, {"_id": 0, "name": 1}).to_list(200)
        oem_vendor_names = {b["name"].lower() for b in oem_brands if b.get("name")}

    # Load catalog with embeddings
    catalog = await db.product_catalog.find(
        {"is_active": True},
        {
            "_id": 0, "article_number": 1, "title_en": 1, "crm_code": 1,
            "vendor": 1, "product_model": 1, "datasheet_url": 1,
            "aliases": 1, "embedding": 1,
        },
    ).limit(MAX_CATALOG_PRODUCTS).to_list(MAX_CATALOG_PRODUCTS)

    if not catalog:
        raise HTTPException(status_code=404, detail="Product catalog is empty")

    # Web research
    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
    logger.info(f"research-item: web research for '{item}'")
    enriched_text, source_urls = _web_research_item(item, client)
    logger.info(f"research-item: enriched='{enriched_text[:80]}' sources={source_urls}")

    # Re-embed enriched description → TOP candidates
    catalog_embeddings = [p.get("embedding") for p in catalog]
    try:
        candidates_list = _voyage_top_k([enriched_text], catalog, catalog_embeddings, PHASE3_TOP_K)
        candidates = candidates_list[0] if candidates_list else []
    except Exception as exc:
        logger.warning(f"research-item Voyage failed: {exc}")
        candidates = _text_fallback_top_k(enriched_text, catalog, PHASE3_TOP_K)

    # Claude re-match
    try:
        claude_results = _claude_match_with_candidates([(item, candidates)])
        r = claude_results[0] if claude_results else {}
    except Exception as exc:
        logger.error(f"research-item Claude failed: {exc}")
        r = {}

    # Normalise confidence
    raw_conf = (r.get("confidence") or "").lower()
    if raw_conf in ("high", "medium", "low"):
        confidence = raw_conf
    else:
        has_match = bool(r.get("matched_title", "").strip())
        confidence = ("medium" if r.get("comment") else "high") if has_match else None

    # Build comment with sources
    comment = r.get("comment") or ""
    if source_urls and bool(r.get("matched_title", "").strip()):
        source_str = " · ".join(source_urls[:3])
        comment = f"Matched via web research. Sources: {source_str}" + (f" | {comment}" if comment else "")

    # Pick code based on mode
    vendor_lower = (r.get("vendor") or "").lower()
    article = r.get("article_number") or ""
    crm = r.get("crm_code") or ""

    if mode == "global":
        code = crm or article
    else:
        code = (article or crm) if (vendor_lower and vendor_lower in oem_vendor_names) else (crm or article)

    return {
        "customer_item": item,
        "crm_code": crm or None,
        "article_number": article or None,
        "matched_title": r.get("matched_title", ""),
        "code": code or None,
        "datasheet_url": r.get("datasheet_url", "") or None,
        "comment": comment or None,
        "web_sources": source_urls,
        "match_type": "web_research",
        "confidence": confidence,
    }


@router.delete("/aliases/{alias_id}")
async def delete_alias(alias_id: str, current_user: dict = Depends(get_current_user)):
    """
    Delete a learned alias from product_aliases collection.
    Requires isAdmin or canEditProductCatalog permission.
    """
    db = get_db()

    is_admin_user = current_user.get("isAdmin") or current_user.get("email", "").endswith("@admin.com")
    can_edit = current_user.get("canEditProductCatalog", False)
    if not is_admin_user and not can_edit:
        raise HTTPException(status_code=403, detail="Permission denied")

    try:
        oid = ObjectId(alias_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid alias id")

    result = await db.product_aliases.delete_one({"_id": oid})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Alias not found")

    return {"message": "Alias deleted"}


@router.post("/generate")
async def generate_excel(
    data: GenerateRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Accept reviewed results, save aliases to product_aliases collection,
    generate and return Excel file.
    """
    db = get_db()
    results = data.results
    mode = data.mode

    # Save aliases (skip duplicates)
    now = datetime.now(timezone.utc).isoformat()
    for r in results:
        customer_item = (r.get("customer_item") or "").strip()
        crm_code = r.get("crm_code") or None
        article_number = r.get("article_number") or None
        if not customer_item:
            continue
        if not crm_code and not article_number:
            continue

        # Check for existing alias
        existing = await db.product_aliases.find_one({
            "alias": customer_item,
            **({"crm_code": crm_code} if crm_code else {"article_number": article_number}),
        })
        if not existing:
            confidence = "confirmed" if r.get("match_type") == "confirmed" else "auto"
            alias_doc = {
                "crm_code": crm_code,
                "article_number": article_number,
                "alias": customer_item,
                "confidence": confidence,
                "saved_at": now,
            }
            await db.product_aliases.insert_one(alias_doc)

    # Build Excel rows from results
    rows = [
        {
            "customer_item": r.get("customer_item", ""),
            "code": r.get("code") or "",
            "matched_title": r.get("matched_title") or "",
            "datasheet_url": r.get("datasheet_url") or "",
            "comment": r.get("comment") or "",
        }
        for r in results
    ]

    try:
        excel_path = _build_excel(rows, mode)
    except Exception as exc:
        logger.error(f"Excel build error: {exc}")
        raise HTTPException(status_code=500, detail=f"Failed to build Excel: {str(exc)[:100]}")

    return FileResponse(
        path=excel_path,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename="product_matching_results.xlsx",
        background=BackgroundTask(_remove_file, excel_path),
    )


@router.get("/search")
async def search_products(
    q: str = Query(..., min_length=1),
    current_user: dict = Depends(get_current_user),
):
    """
    Search product_catalog by title, article_number, aliases.
    Also searches product_aliases collection.
    Returns max 10 deduplicated results.
    """
    db = get_db()
    q_strip = q.strip()
    regex = {"$regex": q_strip, "$options": "i"}

    # Search catalog
    catalog_results = await db.product_catalog.find(
        {
            "is_active": True,
            "$or": [
                {"title_en": regex},
                {"article_number": regex},
                {"aliases": regex},
            ],
        },
        {"_id": 0, "crm_code": 1, "article_number": 1, "title_en": 1, "vendor": 1},
    ).limit(10).to_list(10)

    seen_crm: set = set()
    results = []
    for p in catalog_results:
        key = p.get("crm_code") or p.get("article_number") or ""
        if key and key not in seen_crm:
            seen_crm.add(key)
            results.append({
                "crm_code": p.get("crm_code"),
                "article_number": p.get("article_number"),
                "title": p.get("title_en", ""),
                "code": p.get("crm_code") or p.get("article_number"),
                "vendor": p.get("vendor", ""),
            })

    # Search product_aliases if we have room
    if len(results) < 10:
        alias_hits = await db.product_aliases.find(
            {"alias": regex},
            {"_id": 0, "crm_code": 1, "article_number": 1},
        ).limit(10).to_list(10)

        crm_codes_from_aliases = list({a["crm_code"] for a in alias_hits if a.get("crm_code")} - seen_crm)
        article_nums_from_aliases = list({a["article_number"] for a in alias_hits if a.get("article_number") and not a.get("crm_code")} - seen_crm)

        if crm_codes_from_aliases or article_nums_from_aliases:
            extra = await db.product_catalog.find(
                {
                    "is_active": True,
                    "$or": [
                        *({"crm_code": c} for c in crm_codes_from_aliases),
                        *({"article_number": a} for a in article_nums_from_aliases),
                    ],
                },
                {"_id": 0, "crm_code": 1, "article_number": 1, "title_en": 1, "vendor": 1},
            ).limit(10 - len(results)).to_list(10)

            for p in extra:
                key = p.get("crm_code") or p.get("article_number") or ""
                if key and key not in seen_crm:
                    seen_crm.add(key)
                    results.append({
                        "crm_code": p.get("crm_code"),
                        "article_number": p.get("article_number"),
                        "title": p.get("title_en", ""),
                        "code": p.get("crm_code") or p.get("article_number"),
                        "vendor": p.get("vendor", ""),
                    })

    return results[:10]
